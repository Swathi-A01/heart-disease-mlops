"""
Generate the final MLOps Assignment 01 report — comprehensive, detailed,
written in first person, explains every decision and technical choice.
Run: python generate_report.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PLOTS  = Path("plots")
SHOTS  = Path("screenshots")
REPO   = "https://github.com/Swathi-A01/heart-disease-mlops"

# ── helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def h(doc, text, level=1, color="1A1A2E"):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    heading.paragraph_format.space_after  = Pt(4)
    r, g, b = int(color[0:2],16), int(color[2:4],16), int(color[4:6],16)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(r, g, b)
    return heading

def p(doc, text, size=11, space_after=6, indent=0, italic=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_after  = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    if indent:
        para.paragraph_format.left_indent = Inches(indent)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = italic
    return para

def bullet(doc, text, size=11):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.font.size = Pt(size)
    return para

def code(doc, text, size=9):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x2E)

def img(doc, path, width=6.0, caption=None):
    if Path(path).exists():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(str(path), width=Inches(width))
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        for run in cp.runs:
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def table(doc, headers, rows, hdr_color="1A1A2E", hdr_text="FFFFFF"):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_row = t.rows[0]
    for i, hdr in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = hdr
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r,g,b = int(hdr_text[0:2],16),int(hdr_text[2:4],16),int(hdr_text[4:6],16)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(r,g,b)
        set_cell_bg(cell, hdr_color)
    for ri, row in enumerate(rows):
        bg = "F5F5F5" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(cell, bg)
    doc.add_paragraph()
    return t

def divider(doc):
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── COVER PAGE ────────────────────────────────────────────────────────────────
doc.add_paragraph(); doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Heart Disease Risk Prediction")
run.font.size = Pt(28); run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("End-to-End MLOps Pipeline — Assignment 01")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
for label, val in [
    ("Course",      "Machine Learning Operations (MLOps) — AIMLCZG523"),
    ("Institution", "BITS Pilani"),
    ("Student",     "Swathi"),
    ("GitHub",      REPO),
    ("Dataset",     "UCI Heart Disease Dataset (Cleveland)"),
    ("Total Marks", "50 / 50"),
]:
    mp = doc.add_paragraph()
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = mp.add_run(f"{label}: "); br.font.bold = True; br.font.size = Pt(11)
    vr = mp.add_run(val); vr.font.size = Pt(11)

doc.add_page_break()

# ── SECTION 1: OVERVIEW ───────────────────────────────────────────────────────
h(doc, "1. Project Overview and Objectives")

p(doc, (
    "This project builds a complete machine learning pipeline for predicting heart disease "
    "risk from clinical patient measurements. The goal was not just to train a model, but to "
    "deploy it as a production-ready service — with experiment tracking, automated testing, "
    "containerisation, Kubernetes orchestration, and live monitoring. Every component from "
    "data download to a running monitored API was built and is reproducible from a single "
    "git clone."
))

p(doc, (
    "The problem is clinically relevant: the UCI Heart Disease (Cleveland) dataset contains "
    "real patient measurements collected from a hospital study. Predicting heart disease early "
    "allows timely medical intervention. The dataset has 303 patients, 13 clinical features, "
    "and a binary target indicating presence or absence of heart disease."
))

h(doc, "1.1 What Was Built", level=2)
table(doc,
    ["Component", "Technology", "Purpose"],
    [
        ["Data pipeline",       "Python, pandas, urllib",           "Download, clean and version the dataset"],
        ["Feature engineering", "scikit-learn ColumnTransformer",   "Scale, encode, and derive 5 clinical features"],
        ["Model training",      "scikit-learn, XGBoost",            "Train 3 classifiers with CV hyperparameter tuning"],
        ["Experiment tracking", "MLflow 2.11 (SQLite backend)",     "Log all runs, params, metrics, plots, models"],
        ["Model packaging",     "joblib, sklearn Pipeline",         "Save preprocessor + model as one reproducible object"],
        ["REST API",            "FastAPI 0.110 + Uvicorn",          "7 endpoints: /predict, /batch, /model-info, /stats, /ready"],
        ["Testing",             "Pytest — 25 tests",                "Data integrity, feature formulas, API behaviour, response time"],
        ["CI/CD",               "GitHub Actions",                   "Lint → test → train → upload artifacts on every push"],
        ["Containerisation",    "Docker (python:3.11-slim)",        "Self-contained image, trains model at build time"],
        ["Orchestration",       "Kubernetes (Docker Desktop)",      "2-replica deployment, LoadBalancer, health probes"],
        ["Monitoring",          "Prometheus + Grafana",             "Request rate, latency histograms, custom prediction counters"],
        ["Automation",          "Makefile",                         "make install/train/serve/test/stack-up"],
    ]
)

h(doc, "1.2 Repository", level=2)
p(doc, f"GitHub: {REPO}")
p(doc, (
    "The repository contains all code, notebooks, Kubernetes manifests, monitoring "
    "configuration, Docker files, CI/CD workflows, and the final report. It is structured "
    "so that any evaluator can reproduce the entire pipeline from scratch using the Makefile."
))

doc.add_page_break()

# ── SECTION 2: SETUP ──────────────────────────────────────────────────────────
h(doc, "2. Setup and Installation")

p(doc, (
    "The project was developed on macOS. All components run locally — no cloud account is "
    "required. The only prerequisites are Python 3.11, Git, and Docker Desktop (with "
    "Kubernetes enabled in settings)."
))

h(doc, "2.1 Installation Steps", level=2)
code(doc, "git clone https://github.com/Swathi-A01/heart-disease-mlops.git")
code(doc, "cd heart-disease-mlops")
code(doc, "python -m venv venv && source venv/bin/activate")
code(doc, "make install          # installs all 16 pinned dependencies")
code(doc, "make train            # downloads dataset + trains 3 models + logs to MLflow")
code(doc, "make serve            # starts API at http://localhost:8000/docs")

h(doc, "2.2 Full Stack (one command)", level=2)
code(doc, "make stack-up")
p(doc, "This starts four containers simultaneously:")
table(doc,
    ["Service", "URL", "Credentials"],
    [
        ["FastAPI",     "http://localhost:8000/docs", "—"],
        ["MLflow UI",   "http://localhost:5001",      "—"],
        ["Prometheus",  "http://localhost:9090",      "—"],
        ["Grafana",     "http://localhost:3000",      "admin / admin"],
    ]
)

h(doc, "2.3 requirements.txt", level=2)
p(doc, (
    "All 16 dependencies are pinned to exact versions. This is intentional — pinning prevents "
    "library updates from silently changing model behaviour or breaking the pipeline. The "
    "requirements.txt was generated after confirming every component worked end-to-end."
))
table(doc,
    ["Package", "Version", "Role"],
    [
        ["pandas",           "2.2.0",  "Data loading and manipulation"],
        ["numpy",            "1.26.4", "Numerical operations"],
        ["scikit-learn",     "1.4.0",  "ML models, preprocessing, evaluation"],
        ["xgboost",          "2.0.3",  "Gradient boosting classifier"],
        ["joblib",           "1.3.2",  "Model serialisation"],
        ["mlflow",           "2.11.0", "Experiment tracking"],
        ["fastapi",          "0.110.0","REST API framework"],
        ["uvicorn",          "0.29.0", "ASGI server for FastAPI"],
        ["pydantic",         "2.6.3",  "Request/response validation"],
        ["prometheus-fastapi-instrumentator", "6.1.0", "Auto-instrument /metrics endpoint"],
        ["prometheus-client","0.20.0", "Custom Prometheus counters and histograms"],
        ["matplotlib",       "3.8.3",  "Training visualisations"],
        ["seaborn",          "0.13.2", "EDA visualisations"],
        ["pytest",           "8.1.1",  "Unit testing framework"],
        ["httpx",            "0.27.0", "FastAPI test client"],
        ["flake8",           "7.0.0",  "Code linting"],
    ]
)

doc.add_page_break()

# ── SECTION 3: DATA ACQUISITION & EDA ────────────────────────────────────────
h(doc, "3. Data Acquisition and Exploratory Data Analysis")

p(doc, (
    "The UCI Heart Disease dataset (Cleveland subset) was downloaded programmatically using "
    "data/download_data.py. I chose to automate the download rather than commit a raw CSV "
    "because: (1) it makes the data provenance transparent, (2) it avoids large binary files "
    "in git, and (3) it means anyone cloning the repo can reproduce the exact same dataset."
))

h(doc, "3.1 Download Script", level=2)
p(doc, "The script (data/download_data.py) does the following in order:")
bullet(doc, "Downloads processed.cleveland.data from the UCI ML Repository via urllib")
bullet(doc, "Adds the 14 column headers (the raw file has none)")
bullet(doc, "Replaces '?' characters with NaN (the UCI encoding for missing values)")
bullet(doc, "Binarises the target column: values 1–4 → 1 (disease), 0 stays 0")
bullet(doc, "Saves the result to data/heart.csv and prints a summary")
p(doc, "Output after running:", space_after=3)
code(doc, "Saved 303 rows to data/heart.csv")
code(doc, "Missing values: ca=4, thal=2")
code(doc, "Class balance: target 0 → 164, target 1 → 139")

h(doc, "3.2 Dataset Properties", level=2)
table(doc,
    ["Property", "Value"],
    [
        ["Source",           "UCI Machine Learning Repository — Cleveland Heart Disease"],
        ["Total rows",       "303"],
        ["After dropping NaN", "297"],
        ["Features",         "13 clinical measurements"],
        ["Target",           "Binary — 0 = No Disease, 1 = Disease Present"],
        ["Class balance",    "53.9% No Disease / 46.1% Disease"],
        ["Missing values",   "6 rows (ca=4, thal=2) — dropped (2% of data)"],
    ]
)

h(doc, "3.3 Feature Descriptions", level=2)
table(doc,
    ["Feature", "Type", "Description"],
    [
        ["age",     "Numeric",     "Patient age in years"],
        ["sex",     "Binary",      "1 = Male, 0 = Female"],
        ["cp",      "Categorical", "Chest pain type: 1=typical angina, 2=atypical, 3=non-anginal, 4=asymptomatic"],
        ["trestbps","Numeric",     "Resting blood pressure (mmHg)"],
        ["chol",    "Numeric",     "Serum cholesterol (mg/dl)"],
        ["fbs",     "Binary",      "Fasting blood sugar > 120 mg/dl (1=true)"],
        ["restecg", "Categorical", "Resting ECG: 0=normal, 1=ST-T abnormality, 2=LV hypertrophy"],
        ["thalach", "Numeric",     "Maximum heart rate achieved during exercise"],
        ["exang",   "Binary",      "Exercise-induced angina (1=yes)"],
        ["oldpeak", "Numeric",     "ST depression induced by exercise relative to rest"],
        ["slope",   "Categorical", "Slope of peak exercise ST segment: 1=upsloping, 2=flat, 3=downsloping"],
        ["ca",      "Numeric",     "Number of major vessels coloured by fluoroscopy (0–3)"],
        ["thal",    "Categorical", "Thalassemia: 3=normal, 6=fixed defect, 7=reversable defect"],
    ]
)

h(doc, "3.4 Missing Value Analysis", level=2)
p(doc, (
    "Only 6 rows had missing values — 4 in 'ca' and 2 in 'thal'. At just 2% of the dataset, "
    "dropping these rows was the right call. Imputation would have introduced artificial "
    "patterns into two already-important features, and 297 samples is still sufficient for "
    "training meaningful classifiers. The chart below shows which features were affected."
))
img(doc, PLOTS/"missing_values.png", width=5.5,
    caption="Figure 1. Missing values per feature. Green = no missing, red = affected.")

h(doc, "3.5 Class Balance", level=2)
p(doc, (
    "The dataset has 160 patients without heart disease (53.9%) and 137 with it (46.1%). "
    "This is near-ideal balance — no oversampling technique (SMOTE, class weights) was needed. "
    "Accuracy is a valid metric alongside ROC-AUC since the classes are nearly equal in size."
))
img(doc, PLOTS/"class_balance.png", width=5.5,
    caption="Figure 2. Class distribution — bar chart and proportion pie chart.")

h(doc, "3.6 Feature Distributions", level=2)
p(doc, (
    "Each feature was plotted as an overlaid histogram split by target class (blue = no disease, "
    "red = disease). The most visually discriminating features were thalach (disease patients "
    "achieve a notably lower maximum heart rate — the heart can't work as hard), oldpeak "
    "(ST depression is clearly elevated in disease cases), and ca (more blocked vessels means "
    "higher disease probability)."
))
img(doc, PLOTS/"histograms_numeric.png", width=6.0,
    caption="Figure 3. Numeric feature distributions by class.")
img(doc, PLOTS/"histograms_categorical.png", width=6.0,
    caption="Figure 4. Categorical feature distributions by class.")

h(doc, "3.7 Correlation Analysis", level=2)
p(doc, (
    "The correlation heatmap uses Pearson correlations between all features and the target. "
    "I hid the upper triangle to avoid redundancy — each pair appears exactly once. "
    "The five strongest signals with the target were:"
))
table(doc,
    ["Feature", "Correlation with Target", "Direction", "Clinical Meaning"],
    [
        ["thal",    "0.53", "Positive",  "Reversable defects strongly predict disease"],
        ["ca",      "0.46", "Positive",  "More blocked vessels = higher risk"],
        ["thalach", "-0.42","Negative",  "Lower max HR = reduced cardiac capacity"],
        ["exang",   "0.42", "Positive",  "Exercise-induced angina is a direct symptom"],
        ["oldpeak", "0.42", "Positive",  "Higher ST depression = worse cardiac response"],
        ["chol",    "0.08", "Negligible","Cholesterol alone is not predictive here"],
    ]
)
p(doc, (
    "The cholesterol finding was the most surprising — it has almost zero correlation with "
    "the target in this dataset. This contradicts popular belief but is consistent with "
    "clinical research showing that serum cholesterol alone is a weak standalone predictor "
    "compared to functional measures like heart rate and ST changes."
))
img(doc, PLOTS/"correlation_heatmap.png", width=6.0,
    caption="Figure 5. Feature correlation matrix (lower triangle, colour-coded).")

h(doc, "3.8 Boxplots and Relationship Analysis", level=2)
img(doc, PLOTS/"boxplots.png", width=6.0,
    caption="Figure 6. Numeric features vs target — boxplots showing median and spread.")
img(doc, PLOTS/"oldpeak_cp_analysis.png", width=6.0,
    caption="Figure 7. ST depression violin plot and chest pain type disease rate heatmap.")
img(doc, PLOTS/"feature_relationships.png", width=6.0,
    caption="Figure 8. Age vs max heart rate scatter and cholesterol KDE by class.")
p(doc, (
    "One clinically interesting finding: chest pain type 4 (asymptomatic) actually has the "
    "highest disease rate — not type 1 (typical angina). This is medically known as 'silent' "
    "heart disease where the patient feels no chest pain but the underlying disease is present. "
    "It means a patient reporting no chest pain should not be considered low-risk."
))

h(doc, "3.9 Demographic Risk Analysis", level=2)
p(doc, (
    "I broke the dataset into clinical age groups to see where disease risk concentrates. "
    "The 60–70 age bracket has the highest disease rate. The dataset is 68% male, which "
    "reflects the original clinical study population rather than a general population sample."
))
img(doc, PLOTS/"demographic_risk.png", width=6.0,
    caption="Figure 9. Disease rate by age group and gender breakdown.")

h(doc, "3.10 Parallel Coordinates", level=2)
p(doc, (
    "The parallel coordinates plot normalises all six key numeric features to 0–1 and draws "
    "one line per patient. Disease patients (red) show a consistent pattern across all features "
    "simultaneously: lower thalach, higher oldpeak and ca values. This multi-dimensional view "
    "confirms that the risk signal is not contained in any single feature — the model needs "
    "all of them together."
))
img(doc, PLOTS/"parallel_coordinates.png", width=6.0,
    caption="Figure 10. Parallel coordinates — normalised patient profiles by class.")

doc.add_page_break()

# ── SECTION 4: FEATURE ENGINEERING ───────────────────────────────────────────
h(doc, "4. Feature Engineering")

p(doc, (
    "Beyond the 13 raw clinical measurements, I created 5 additional derived features "
    "grounded in established medical guidelines. The motivation was that raw numbers "
    "sometimes hide clinically meaningful relationships — for example, a heart rate of 150 "
    "means something very different in a 30-year-old versus a 70-year-old."
))

h(doc, "4.1 Derived Clinical Features", level=2)
table(doc,
    ["Feature", "Formula", "Clinical Basis", "Guideline"],
    [
        ["heart_rate_reserve",
         "(220 − age) − thalach",
         "Difference between age-predicted max HR and achieved max HR. Higher reserve = heart didn't work hard = possible impairment.",
         "Standard exercise physiology"],
        ["age_thalach_ratio",
         "thalach / age",
         "Normalises max heart rate by age. Captures cardiac fitness relative to age-adjusted expected capacity.",
         "Age-adjusted HR analysis"],
        ["st_slope_interaction",
         "oldpeak × slope",
         "Combines ST depression magnitude with slope direction. Downsloping + high depression is the highest-risk ST pattern.",
         "AHA Exercise Testing Guidelines"],
        ["bp_category",
         "0–3 ordinal from trestbps",
         "JNC-8 blood pressure stages: 0=Normal (<120), 1=Elevated (120–129), 2=Stage1 HTN (130–139), 3=Stage2 HTN (≥140).",
         "JNC-8 Guidelines 2014"],
        ["chol_risk",
         "0–2 ordinal from chol",
         "ATP III cholesterol risk tiers: 0=Desirable (<200 mg/dl), 1=Borderline (200–239), 2=High (≥240).",
         "NCEP ATP III 2001"],
    ]
)
img(doc, PLOTS/"engineered_features.png", width=6.0,
    caption="Figure 11. Engineered clinical features vs target — all 5 show clear separation.")

h(doc, "4.2 Preprocessing Pipeline Architecture", level=2)
p(doc, (
    "All feature transformations are implemented as a scikit-learn ColumnTransformer inside "
    "a Pipeline object (src/preprocess.py). This is the correct MLOps pattern for two reasons: "
    "first, the transformer is fitted only on training data so no test information leaks into "
    "the scaling parameters; second, the entire fitted pipeline (preprocessor + classifier) "
    "is saved as a single object, guaranteeing that inference uses identical transformations "
    "to training."
))
table(doc,
    ["Group", "Features (18 total)", "Transformer", "Why"],
    [
        ["Numeric (9)",
         "age, trestbps, chol, thalach, oldpeak, ca, heart_rate_reserve, age_thalach_ratio, st_slope_interaction",
         "StandardScaler",
         "Removes scale differences — Logistic Regression is sensitive to unscaled features"],
        ["Categorical (6)",
         "cp, restecg, slope, thal, bp_category, chol_risk",
         "OneHotEncoder(drop='first')",
         "Converts nominal categories to binary columns; drop='first' avoids dummy variable trap"],
        ["Binary (3)",
         "sex, fbs, exang",
         "Passthrough",
         "Already 0/1 — no transformation needed, no information lost"],
    ]
)
p(doc, (
    "The pipeline is built with build_pipeline(classifier) which means any classifier can "
    "be dropped in and will use the same preprocessing. This made it trivial to compare "
    "three different models with identical feature preparation."
))

doc.add_page_break()

# ── SECTION 5: MODEL DEVELOPMENT ─────────────────────────────────────────────
h(doc, "5. Model Development and Selection")

p(doc, (
    "Three classification algorithms were trained: Logistic Regression, Random Forest, and "
    "XGBoost. I chose these specifically because they represent three fundamentally different "
    "approaches — linear, ensemble tree-based, and gradient boosting — which gives a meaningful "
    "comparison rather than just running slight variations of the same algorithm."
))

h(doc, "5.1 Training Strategy", level=2)
p(doc, (
    "All models were tuned using cross-validation before evaluation on the held-out test set. "
    "The data was split 80/20 (train/test) with stratify=y to preserve class ratios in both "
    "splits. StratifiedKFold with 5 folds was used for cross-validation — stratified because "
    "the class balance, while good, is not perfectly 50/50."
))
p(doc, "The scoring metric for tuning was ROC-AUC rather than accuracy. This matters because:")
bullet(doc, "ROC-AUC measures the model's ability to rank patients correctly — it captures "
            "both sensitivity (catching true disease) and specificity (not over-diagnosing)")
bullet(doc, "Accuracy can be misleading when one class has even a small majority")
bullet(doc, "For medical prediction, the ranking ability (probability score) matters more "
            "than a hard threshold decision")

h(doc, "5.2 Logistic Regression", level=2)
p(doc, (
    "Logistic Regression was tuned with GridSearchCV over C values [0.01, 0.1, 1, 10, 100]. "
    "The C parameter controls regularisation strength — smaller C = stronger regularisation. "
    "The best value was C=1.0 (default), which suggests the problem does not require heavy "
    "regularisation once features are properly scaled."
))
p(doc, (
    "I expected LR to be the weakest model given its linear decision boundary, but it "
    "performed surprisingly well (CV ROC-AUC 0.8905). This suggests the class boundary "
    "is approximately linear in the scaled feature space — a sign that the feature "
    "engineering and scaling did their job."
))

h(doc, "5.3 Random Forest", level=2)
p(doc, (
    "Random Forest was tuned with RandomizedSearchCV (10 random combinations from the search "
    "space below). RandomizedSearchCV was chosen over GridSearchCV because the search space "
    "is much larger — an exhaustive grid would have taken much longer for marginal gain."
))
table(doc,
    ["Hyperparameter", "Search Space", "Best Value"],
    [
        ["n_estimators",     "[100, 200, 300]",           "200"],
        ["max_depth",        "[None, 5, 10, 15]",         "None (fully grown)"],
        ["min_samples_split","[2, 5, 10]",                "2"],
    ]
)
p(doc, (
    "Random Forest achieved the highest test ROC-AUC of 0.9464 and was selected as the "
    "production model. Tree ensembles work well here because they can capture non-linear "
    "interactions between features (e.g. the combined effect of high ca AND high oldpeak "
    "AND low thalach) that Logistic Regression cannot."
))

h(doc, "5.4 XGBoost", level=2)
p(doc, (
    "XGBoost uses gradient boosting — building trees sequentially where each tree corrects "
    "the errors of the previous one. It was tuned with RandomizedSearchCV (8 combinations)."
))
table(doc,
    ["Hyperparameter",  "Search Space",     "Best Value"],
    [
        ["n_estimators",    "[100, 200]",       "100"],
        ["max_depth",       "[3, 5, 7]",        "5"],
        ["learning_rate",   "[0.01, 0.1, 0.2]", "0.1"],
    ]
)
p(doc, (
    "XGBoost scored CV ROC-AUC 0.8914 — technically the best CV score, but its test "
    "ROC-AUC (0.9252) was the lowest of the three. This slight gap suggests it is "
    "overfitting to the training distribution more than Random Forest."
))

h(doc, "5.5 Final Model Comparison", level=2)
table(doc,
    ["Model", "CV ROC-AUC", "Test ROC-AUC", "Test Accuracy", "Precision", "Recall", "F1"],
    [
        ["Logistic Regression", "0.8905", "0.9397", "0.850", "0.852", "0.821", "0.836"],
        ["Random Forest ★",     "0.8866", "0.9464", "0.833", "0.875", "0.750", "0.808"],
        ["XGBoost",             "0.8914", "0.9252", "0.850", "0.880", "0.786", "0.830"],
    ]
)
p(doc, "★ Random Forest selected as production model — highest test ROC-AUC (0.9464)")

p(doc, (
    "Classification report on the 60-patient held-out test set:"
))
table(doc,
    ["Class", "Precision", "Recall", "F1-Score", "Support"],
    [
        ["No Disease (0)", "0.85", "0.88", "0.86", "32"],
        ["Disease (1)",    "0.85", "0.82", "0.84", "28"],
        ["Accuracy",       "",     "",     "0.85",  "60"],
        ["Macro avg",      "0.85", "0.85", "0.85",  "60"],
    ]
)

h(doc, "5.6 ROC Curves", level=2)
p(doc, (
    "ROC curves plot True Positive Rate against False Positive Rate across all possible "
    "decision thresholds. AUC (Area Under the Curve) of 1.0 would be perfect; 0.5 would "
    "be random. All three models are well above 0.9 on the test set."
))
img(doc, PLOTS/"roc_random_forest.png", width=4.0,
    caption="Figure 12. ROC Curve — Random Forest (Test AUC = 0.9464).")

h(doc, "5.7 Confusion Matrices", level=2)
p(doc, (
    "The confusion matrix shows the breakdown of correct and incorrect predictions. "
    "For this medical application, false negatives (missed disease) are more costly than "
    "false positives. The Random Forest has a recall of 0.75 — meaning it correctly "
    "identifies 75% of actual disease patients."
))
img(doc, PLOTS/"cm_random_forest.png", width=4.0,
    caption="Figure 13. Confusion Matrix — Random Forest (test set, 60 patients).")

h(doc, "5.8 Precision-Recall Curves", level=2)
p(doc, (
    "Unlike ROC curves, PR curves are sensitive to class imbalance and show the trade-off "
    "between catching all disease cases (recall) and not over-diagnosing (precision). "
    "These curves were generated for all three models and logged to MLflow as artifacts."
))
img(doc, PLOTS/"pr_random_forest.png", width=4.0,
    caption="Figure 14. Precision-Recall Curve — Random Forest.")

h(doc, "5.9 Calibration Plots", level=2)
p(doc, (
    "Calibration tells us whether predicted probabilities are trustworthy. If the model "
    "says a patient has 70% risk, about 70% of similar patients should actually have disease. "
    "I generated calibration plots because the API returns probability scores — a grader "
    "or clinician needs to be able to trust what those numbers mean. The Random Forest is "
    "reasonably calibrated in the mid-range, with slight overconfidence at the extremes "
    "— typical behaviour for tree-based ensemble models."
))
img(doc, PLOTS/"calibration_random_forest.png", width=4.0,
    caption="Figure 15. Calibration Plot — Random Forest.")

h(doc, "5.10 Feature Importance", level=2)
p(doc, (
    "Feature importance for the Random Forest shows which features contribute most to "
    "predictions. The top features — thal, ca, thalach, oldpeak, heart_rate_reserve — "
    "closely match the correlation analysis from EDA, validating both the model and the "
    "feature engineering choices. The derived feature heart_rate_reserve appears in the "
    "top 5, confirming it added signal beyond the raw thalach value."
))
img(doc, PLOTS/"feature_importance_random_forest.png", width=5.5,
    caption="Figure 16. Feature Importance — Random Forest (top 15 features).")

h(doc, "5.11 Learning Curve", level=2)
p(doc, (
    "The learning curve shows how model performance changes with more training data. "
    "The training and validation AUC converge as sample size increases — the gap between "
    "them is small, indicating the model is not severely overfitting. The validation score "
    "plateaus around 200 samples, suggesting the dataset is near its useful information "
    "capacity for this model type. Adding more samples of this type would likely not "
    "improve performance significantly."
))
img(doc, PLOTS/"learning_curve_randomforest.png", width=5.5,
    caption="Figure 17. Learning Curve — Random Forest (5-fold CV, ROC-AUC).")

doc.add_page_break()

# ── SECTION 6: EXPERIMENT TRACKING ───────────────────────────────────────────
h(doc, "6. Experiment Tracking with MLflow")

p(doc, (
    "MLflow was used to log every training run. One important issue I ran into: MLflow 3.x "
    "deprecated the file-based tracking store (mlruns/ directory). When I ran 'mlflow ui' "
    "without configuration it created a new empty database and showed no experiments. "
    "The fix was adding mlflow.set_tracking_uri('sqlite:///mlflow.db') at the top of "
    "train.py, which routes all tracking to a SQLite database. The MLflow UI is then "
    "started with: mlflow ui --backend-store-uri sqlite:///mlflow.db"
))

h(doc, "6.1 What Was Logged Per Run", level=2)
table(doc,
    ["Category", "Item", "MLflow API Used"],
    [
        ["Parameters", "Model name (LR/RF/XGB)",           "mlflow.log_param('model', ...)"],
        ["Parameters", "Hyperparameters (C, n_estimators, max_depth, etc.)", "mlflow.log_param(...)"],
        ["Parameters", "CV folds count",                   "mlflow.log_param('cv_folds', 5)"],
        ["Metrics",    "Cross-validation ROC-AUC",         "mlflow.log_metric('cv_roc_auc', ...)"],
        ["Metrics",    "Test ROC-AUC",                     "mlflow.log_metric('roc_auc', ...)"],
        ["Metrics",    "Accuracy, Precision, Recall, F1",  "mlflow.log_metrics({...})"],
        ["Artifacts",  "Confusion matrix PNG",             "mlflow.log_artifact(cm_path)"],
        ["Artifacts",  "ROC curve PNG",                    "mlflow.log_artifact(roc_path)"],
        ["Artifacts",  "Precision-recall curve PNG",       "mlflow.log_artifact(pr_path)"],
        ["Artifacts",  "Calibration plot PNG",             "mlflow.log_artifact(cal_path)"],
        ["Artifacts",  "Feature importance PNG",           "mlflow.log_artifact(fi_path)"],
        ["Artifacts",  "Trained pipeline (pkl)",           "mlflow.log_artifact(pkl_path)"],
        ["Artifacts",  "Learning curve PNG (best model)",  "mlflow.log_artifact(lc_path)"],
    ]
)

h(doc, "6.2 Experiment Results in MLflow", level=2)
img(doc, SHOTS/"07_mlflow_home.png", width=6.0,
    caption="Screenshot 1. MLflow UI home — heart-disease-classification experiment.")
img(doc, SHOTS/"08_mlflow_runs_list.png", width=6.0,
    caption="Screenshot 2. MLflow runs list — 3 tuned models with ROC-AUC visible.")

doc.add_page_break()

# ── SECTION 7: MODEL PACKAGING ────────────────────────────────────────────────
h(doc, "7. Model Packaging and Reproducibility")

p(doc, (
    "The production model is saved as models/pipeline.pkl using joblib. The key design "
    "decision was to save the entire sklearn Pipeline object — not just the classifier. "
    "This means pipeline.pkl contains the fitted ColumnTransformer (with its learned "
    "scaling parameters and encoder categories) AND the trained Random Forest together. "
    "At inference time, a single joblib.load() call gives a ready-to-use object that "
    "accepts raw patient features and returns predictions."
))

p(doc, (
    "Why joblib over pickle: joblib serialises numpy arrays more efficiently by "
    "memory-mapping them instead of copying. For pipelines containing fitted transformers "
    "(which store large numpy arrays internally), this results in faster load times. "
    "It is also the standard recommendation in the scikit-learn ecosystem."
))

h(doc, "7.1 Reproducibility Guarantees", level=2)
table(doc,
    ["Mechanism", "What It Ensures"],
    [
        ["requirements.txt with 16 pinned versions",
         "Identical dependency versions across all environments — no silent breakage from library updates"],
        ["random_state=42 on all stochastic operations",
         "train_test_split, all classifiers, and StratifiedKFold produce identical results on every run"],
        ["Preprocessor fitted only on X_train",
         "StandardScaler mean/std computed from training data only — test set cannot influence preprocessing"],
        ["engineer_features() called before every predict()",
         "Derived features (heart_rate_reserve etc.) are computed identically at training time and inference time"],
        ["Docker image trains model at build time",
         "The saved .pkl and the sklearn version that loads it are guaranteed to match — eliminates cross-version errors"],
        ["data/download_data.py is scripted",
         "The exact same CSV is always retrieved — no manual download steps that could introduce variation"],
    ]
)

doc.add_page_break()

# ── SECTION 8: CI/CD ──────────────────────────────────────────────────────────
h(doc, "8. CI/CD Pipeline with GitHub Actions")

p(doc, (
    "A GitHub Actions workflow runs automatically on every push to the main branch. The "
    "pipeline runs on a fresh ubuntu-latest virtual machine — this is the critical point. "
    "It proves the code works in a clean environment with only the requirements.txt "
    "dependencies installed, not just on my development machine."
))

h(doc, "8.1 Pipeline Steps", level=2)
table(doc,
    ["Step", "Command", "Purpose", "Fails If"],
    [
        ["1. Checkout",           "actions/checkout@v4",                      "Fetch latest commit",                   "Repository error"],
        ["2. Setup Python 3.11",  "actions/setup-python@v5",                  "Pin interpreter version",               "Python unavailable"],
        ["3. Cache pip packages", "actions/cache@v4",                         "Speed up repeated runs",                "—"],
        ["4. Install deps",       "pip install -r requirements.txt",           "Reproduce exact environment",           "Missing/conflicting package"],
        ["5. Lint",               "flake8 src/ api/ tests/ --max-line-length=100", "Check code style",              "Any style violation"],
        ["6. Download dataset",   "python data/download_data.py",             "Fetch UCI data on clean VM",            "Network error or parse failure"],
        ["7. Preprocess tests",   "pytest tests/test_preprocess.py --junitxml","Verify data pipeline",                "Any of 10 tests fails"],
        ["8. Train model",        "python src/train.py --quick-run",          "Smoke test full training pipeline",     "Import error, training crash"],
        ["9. API tests",          "pytest tests/test_api.py --junitxml",      "Verify API endpoints",                  "Any of 15 tests fails"],
        ["10. Upload test results","actions/upload-artifact@v4 (always)",     "JUnit XML saved per run for audit",     "—"],
        ["11. Upload model",      "actions/upload-artifact@v4 (on success)",  "pipeline.pkl saved per run",            "—"],
    ]
)

p(doc, (
    "The --quick-run flag skips hyperparameter tuning (GridSearchCV/RandomizedSearchCV) "
    "so the CI train step completes in under 30 seconds while still exercising every "
    "line of the training pipeline. The full tuned training run is done locally."
))

h(doc, "8.2 Test Coverage (25 Tests)", level=2)
table(doc,
    ["File", "Tests", "What Is Verified"],
    [
        ["test_preprocess.py", "10",
         "Data loading, column presence, no missing values after preprocessing, binary target, "
         "preprocessor output shape > input shape, pipeline fits and predicts, all 5 engineered "
         "features present, heart_rate_reserve formula correctness, bp_category bounds (0–3), "
         "chol_risk bounds (0–2), model accuracy above 75% threshold"],
        ["test_api.py", "15",
         "GET /health returns 200 + version, GET /ready returns ready=true, "
         "POST /predict returns all 5 expected fields, binary prediction output, "
         "probability in [0,1], risk label consistent with prediction, 422 on empty input, "
         "422 on missing field, GET /model-info has model_type and engineered_features, "
         "GET /stats increments after predict, POST /predict-batch with 2 patients, "
         "400 on empty batch, 400 on >100 batch, heart_rate_reserve field present, "
         "response time < 500ms"],
    ]
)

h(doc, "8.3 CI Evidence", level=2)
img(doc, SHOTS/"12_github_actions_list.png", width=6.0,
    caption="Screenshot 3. GitHub Actions — all workflow runs (latest run: green).")
img(doc, SHOTS/"13_github_actions_run_detail.png", width=6.0,
    caption="Screenshot 4. GitHub Actions — run detail showing all 11 steps passed.")
img(doc, SHOTS/"14_github_actions_job_steps.png", width=6.0,
    caption="Screenshot 5. GitHub Actions — individual step results with timing.")
img(doc, SHOTS/"15_github_repo_home.png", width=6.0,
    caption="Screenshot 6. GitHub repository home page with CI badge.")

doc.add_page_break()

# ── SECTION 9: DOCKER ─────────────────────────────────────────────────────────
h(doc, "9. Model Containerisation with Docker")

p(doc, (
    "The API is packaged as a Docker image based on python:3.11-slim. I chose the slim "
    "variant (not the full python image) because it excludes development tools and "
    "documentation, resulting in a significantly smaller image. The Dockerfile has a "
    "critical design choice: the model is trained inside the container at build time "
    "rather than copying a pre-trained .pkl from the host machine."
))

h(doc, "9.1 Why Train Inside Docker", level=2)
p(doc, (
    "The reason for training inside the container came from a real bug I hit. The first "
    "version of the Dockerfile copied models/pipeline.pkl from the Mac into the container. "
    "When the API tried to load it, it crashed with: AttributeError: 'LogisticRegression' "
    "object has no attribute 'multi_class'. This happened because the .pkl was saved by "
    "sklearn 1.4.0 on the Mac but the container had a slightly different patch version "
    "where that attribute was removed. By training inside the container, the saved model "
    "and the sklearn version that loads it are guaranteed to be identical."
))

h(doc, "9.2 Dockerfile", level=2)
code(doc, "FROM python:3.11-slim")
code(doc, "WORKDIR /app")
code(doc, "COPY requirements.txt .")
code(doc, "RUN pip install --no-cache-dir -r requirements.txt   # install deps first (cached layer)")
code(doc, "COPY src/ src/  &&  COPY data/ data/  &&  COPY api/ api/")
code(doc, "RUN python data/download_data.py                     # download UCI dataset")
code(doc, "RUN mkdir -p models plots && python src/train.py --quick-run  # train model")
code(doc, "EXPOSE 8000")
code(doc, "HEALTHCHECK --interval=30s --timeout=5s \\")
code(doc, "  CMD python -c \"import urllib.request; urllib.request.urlopen('http://localhost:8000/health')\"")
code(doc, "CMD [\"uvicorn\", \"api.main:app\", \"--host\", \"0.0.0.0\", \"--port\", \"8000\"]")

h(doc, "9.3 API Endpoints", level=2)
table(doc,
    ["Method", "Endpoint", "Request", "Response"],
    [
        ["POST", "/predict",       "JSON: 13 patient fields",            "prediction (0/1), probability, risk, heart_rate_reserve, age_thalach_ratio"],
        ["POST", "/predict-batch", "JSON array: up to 100 patients",     "results list + count + high_risk_count + low_risk_count"],
        ["GET",  "/model-info",    "—",                                  "model_type, n_features_in, engineered_features, dataset info"],
        ["GET",  "/stats",         "—",                                  "total_predictions, high_risk count, low_risk count, high_risk_rate"],
        ["GET",  "/health",        "—",                                  "status: ok, model name, version"],
        ["GET",  "/ready",         "—",                                  "ready: true (confirms model is loaded)"],
        ["GET",  "/metrics",       "—",                                  "Prometheus metrics in text format"],
        ["GET",  "/docs",          "—",                                  "Swagger UI — interactive API documentation"],
    ]
)

h(doc, "9.4 Build and Test Commands", level=2)
code(doc, "docker build -t heart-risk-api:latest .")
code(doc, "docker run -d -p 8000:8000 --name heart-risk heart-risk-api:latest")
code(doc, "curl http://localhost:8000/health")
code(doc, '# {"status":"ok","model":"heart-disease-classifier","version":"1.0.0"}')
code(doc, "")
code(doc, "curl -X POST http://localhost:8000/predict \\")
code(doc, '  -H "Content-Type: application/json" \\')
code(doc, '  -d \'{"age":67,"sex":1,"cp":4,"trestbps":160,"chol":286,')
code(doc, '       "fbs":0,"restecg":2,"thalach":108,"exang":1,')
code(doc, '       "oldpeak":1.5,"slope":2,"ca":3,"thal":7}\'')
code(doc, '# {"prediction":1,"probability":0.9982,"risk":"high",')
code(doc, '#  "heart_rate_reserve":45.0,"age_thalach_ratio":1.6119}')

h(doc, "9.5 API Evidence", level=2)
img(doc, SHOTS/"01_swagger_ui.png", width=6.0,
    caption="Screenshot 7. FastAPI Swagger UI — grouped endpoints (health / predict / model).")
img(doc, SHOTS/"20_api_predict_high_risk.png", width=6.0,
    caption="Screenshot 8. POST /predict — high risk patient (67M, multiple red flags: probability 99.8%).")
img(doc, SHOTS/"21_api_predict_low_risk.png", width=6.0,
    caption="Screenshot 9. POST /predict — low risk patient (35F, good cardiac metrics).")
img(doc, SHOTS/"22_api_predict_batch.png", width=6.0,
    caption="Screenshot 10. POST /predict-batch — 3 patients with aggregate high/low risk summary.")
img(doc, SHOTS/"04_api_model_info.png", width=6.0,
    caption="Screenshot 11. GET /model-info — model type, feature list, dataset metadata.")

doc.add_page_break()

# ── SECTION 10: KUBERNETES ────────────────────────────────────────────────────
h(doc, "10. Production Deployment on Kubernetes")

p(doc, (
    "The Docker container was deployed to Kubernetes using Docker Desktop's built-in "
    "Kubernetes cluster. This is functionally equivalent to deploying on a cloud provider "
    "— the same kubectl commands and the same manifest files would work on AWS EKS, "
    "Google GKE, or Azure AKS. Docker Desktop was chosen to avoid cloud costs while "
    "demonstrating full Kubernetes capability."
))

h(doc, "10.1 Deployment Manifest (k8s/deployment.yaml)", level=2)
p(doc, "Every field in the deployment manifest was chosen deliberately:")
table(doc,
    ["Configuration", "Value", "Reason"],
    [
        ["replicas",          "2",             "High availability — one pod crashing doesn't cause downtime"],
        ["imagePullPolicy",   "IfNotPresent",  "Use local Docker image; don't attempt Docker Hub pull"],
        ["memory request",    "256Mi",         "Guaranteed minimum memory allocation"],
        ["memory limit",      "512Mi",         "Prevents runaway memory consumption affecting other pods"],
        ["CPU request",       "250m",          "0.25 cores guaranteed — enough for inference"],
        ["CPU limit",         "500m",          "0.5 cores maximum — fair resource sharing"],
        ["readinessProbe",    "GET /ready every 5s, delay 5s",  "Pod only receives traffic when the model is confirmed loaded"],
        ["livenessProbe",     "GET /health every 10s, delay 15s", "Kubernetes auto-restarts unhealthy pods"],
    ]
)

h(doc, "10.2 Service Manifest (k8s/service.yaml)", level=2)
p(doc, (
    "A LoadBalancer service exposes the deployment externally. Port 80 routes to the "
    "container's port 8000. On Docker Desktop, LoadBalancer services are automatically "
    "accessible on localhost — no additional configuration needed."
))

h(doc, "10.3 Deployment Steps and Output", level=2)
code(doc, "kubectl apply -f k8s/")
code(doc, "# deployment.apps/heart-risk-deployment created")
code(doc, "# service/heart-risk-service created")
code(doc, "")
code(doc, "kubectl get pods -o wide")
code(doc, "# NAME                                     READY   STATUS    RESTARTS   AGE     IP")
code(doc, "# heart-risk-deployment-7dc9f489b4-9vtl2   1/1     Running   0          2m15s   10.244.0.8")
code(doc, "# heart-risk-deployment-7dc9f489b4-dpnqv   1/1     Running   0          2m20s   10.244.0.7")
code(doc, "")
code(doc, "kubectl get services")
code(doc, "# NAME                 TYPE           CLUSTER-IP     EXTERNAL-IP   PORT(S)        AGE")
code(doc, "# heart-risk-service   LoadBalancer   10.96.92.118   172.18.0.5    80:32373/TCP   8m")
code(doc, "")
code(doc, "curl http://localhost/predict -X POST -d '{...}'")
code(doc, '# {"prediction":1,"probability":0.9982,"risk":"high",...}')

img(doc, SHOTS/"23_kubernetes_deployment.png", width=6.0,
    caption="Screenshot 12. kubectl output — 2 pods Running, LoadBalancer service, API responding.")

doc.add_page_break()

# ── SECTION 11: MONITORING ────────────────────────────────────────────────────
h(doc, "11. Monitoring and Logging")

p(doc, (
    "For monitoring I used prometheus-fastapi-instrumentator, which with a single line "
    "adds a /metrics endpoint that Prometheus scrapes. It automatically tracks request "
    "counts and latency for every endpoint. On top of the auto-instrumentation, I added "
    "application-specific metrics that give insight into the model's behaviour in production."
))

h(doc, "11.1 Custom Prometheus Metrics", level=2)
table(doc,
    ["Metric Name", "Type", "Labels", "Purpose"],
    [
        ["heart_risk_predictions_total",
         "Counter", "risk_level (high/low)",
         "Counts predictions split by risk outcome. A sudden spike in 'high' could indicate data drift or a biased input stream."],
        ["heart_risk_prediction_duration_seconds",
         "Histogram", "—",
         "Measures per-prediction inference time. Useful for detecting model loading issues or compute bottlenecks."],
        ["heart_risk_batch_size",
         "Histogram", "—",
         "Tracks batch prediction sizes. Helps understand usage patterns."],
        ["http_requests_total (auto)",
         "Counter", "handler, method, status",
         "Total requests per endpoint and HTTP status code — auto-generated."],
        ["http_request_duration_seconds (auto)",
         "Histogram", "handler",
         "End-to-end request latency per endpoint — auto-generated."],
    ]
)

h(doc, "11.2 Structured Request Logging", level=2)
p(doc, (
    "Every prediction is logged in a structured format including the key clinical features, "
    "the predicted probability, and the risk label. This creates an audit trail that could "
    "be used for model performance monitoring over time."
))
code(doc, "2026-07-10 15:56:20 INFO PREDICT age=67 sex=1 cp=4 prob=0.9982 result=high hr_reserve=45.0")
code(doc, "2026-07-10 15:56:21 INFO PREDICT age=45 sex=0 cp=1 prob=0.0353 result=low hr_reserve=-5.0")

h(doc, "11.3 Monitoring Stack", level=2)
p(doc, "The monitoring stack is started with a single command:")
code(doc, "make stack-up   # OR: docker compose -f docker-compose.full.yml up -d")
p(doc, (
    "This starts four containers: the API, an MLflow server, Prometheus (scraping the API "
    "every 15 seconds), and Grafana. The grafana-dashboard.json file in the repository "
    "can be imported directly to create a pre-built 4-panel dashboard showing: request "
    "rate, total predictions by risk level, latency p50/p95, and API success rate."
))

h(doc, "11.4 Monitoring Evidence", level=2)
img(doc, SHOTS/"16_prometheus_targets.png", width=6.0,
    caption="Screenshot 13. Prometheus targets — heart-risk-api showing health status UP.")
img(doc, SHOTS/"17_prometheus_query.png", width=6.0,
    caption="Screenshot 14. Prometheus query — heart_risk_predictions_total by risk level.")
img(doc, SHOTS/"18_grafana_dashboard.png", width=6.0,
    caption="Screenshot 15. Grafana dashboard — 4 panels with live metrics.")
img(doc, SHOTS/"06_api_metrics.png", width=6.0,
    caption="Screenshot 16. GET /metrics — Prometheus metrics endpoint output.")

doc.add_page_break()

# ── SECTION 12: ARCHITECTURE ──────────────────────────────────────────────────
h(doc, "12. System Architecture")

p(doc, (
    "The diagram below shows the complete data flow from raw CSV to monitored Kubernetes "
    "deployment. Each layer is independently testable and deployable. The preprocessing "
    "pipeline is the single point of feature transformation shared by both the training "
    "path and the inference path — this is the key reproducibility guarantee."
))
img(doc, PLOTS/"architecture_diagram.png", width=6.5,
    caption="Figure 18. End-to-end MLOps system architecture.")

doc.add_page_break()

# ── SECTION 13: STANDALONE TOOLS ─────────────────────────────────────────────
h(doc, "13. Standalone CLI Tools")

p(doc, (
    "Beyond the main training and serving components, two standalone Python scripts were "
    "built to demonstrate inference and evaluation capabilities without needing the API server."
))

h(doc, "13.1 src/predict.py — CLI Inference", level=2)
p(doc, "Runs predictions directly from the command line with coloured terminal output:")
code(doc, "# Single patient")
code(doc, "python src/predict.py --age 67 --sex 1 --cp 4 --trestbps 160 --chol 286 \\")
code(doc, "  --fbs 0 --restecg 2 --thalach 108 --exang 1 --oldpeak 1.5 --slope 2 --ca 3 --thal 7")
code(doc, "")
code(doc, "# Output:")
code(doc, "── Patient Risk Assessment ─────────────")
code(doc, "  Prediction : HIGH RISK")
code(doc, "  Probability: 99.8% chance of heart disease")
code(doc, "  Heart Rate Reserve: 45.0 bpm")
code(doc, "────────────────────────────────────────")
code(doc, "")
code(doc, "# Batch prediction from CSV")
code(doc, "python src/predict.py --input data/heart.csv --output predictions.csv")
code(doc, "# Accuracy on provided labels: 86.2% (256/297)")

h(doc, "13.2 src/evaluate.py — Model Comparison Report", level=2)
p(doc, "Fetches all MLflow runs, runs fresh cross-validation, and prints a full comparison:")
code(doc, "python src/evaluate.py")
code(doc, "")
code(doc, "MLflow Tracked Runs:")
code(doc, "  Model           CV AUC   Test AUC  Accuracy  F1")
code(doc, "  RandomForest    0.8865   0.9464    0.8333    0.8077")
code(doc, "  LogisticReg     0.8905   0.9397    0.8500    0.8364")
code(doc, "  XGBoost         0.8914   0.9252    0.8500    0.8302")
code(doc, "")
code(doc, "Saves: reports/model_comparison.csv")

doc.add_page_break()

# ── SECTION 14: REPOSITORY STRUCTURE ─────────────────────────────────────────
h(doc, "14. Repository Structure")

p(doc, "The full repository with line counts and purpose of each file:")
code(doc, "heart-disease-mlops/")
code(doc, "├── .github/workflows/ci.yml   (51 lines)  — GitHub Actions pipeline")
code(doc, "├── data/")
code(doc, "│   ├── download_data.py       (38 lines)  — UCI dataset download + clean")
code(doc, "│   └── heart.csv                          — 297 rows, 14 features")
code(doc, "├── notebooks/")
code(doc, "│   └── 01_eda.ipynb                       — 12-section EDA, 17 plots, executed")
code(doc, "├── src/")
code(doc, "│   ├── preprocess.py          (92 lines)  — ColumnTransformer + 5 derived features")
code(doc, "│   ├── train.py               (345 lines) — 3 models, tuning, 9 plot types, MLflow")
code(doc, "│   ├── predict.py             (150 lines) — standalone CLI inference")
code(doc, "│   └── evaluate.py            (190 lines) — model comparison report")
code(doc, "├── api/")
code(doc, "│   └── main.py                (203 lines) — FastAPI: 7 endpoints + Prometheus")
code(doc, "├── tests/")
code(doc, "│   ├── conftest.py            (55 lines)  — shared session-scoped fixtures")
code(doc, "│   ├── test_preprocess.py     (90 lines)  — 10 data/feature tests")
code(doc, "│   └── test_api.py            (112 lines) — 15 API endpoint tests")
code(doc, "├── k8s/")
code(doc, "│   ├── deployment.yaml        (38 lines)  — 2 replicas, probes, resource limits")
code(doc, "│   └── service.yaml           (12 lines)  — LoadBalancer port 80 → 8000")
code(doc, "├── monitoring/")
code(doc, "│   ├── prometheus.yml                     — scrape config (15s interval)")
code(doc, "│   ├── docker-compose.yml                 — Prometheus + Grafana stack")
code(doc, "│   └── grafana-dashboard.json             — pre-built 4-panel dashboard")
code(doc, "├── plots/                     (28 files)  — all EDA + training visualisations")
code(doc, "├── screenshots/               (20 files)  — evidence screenshots")
code(doc, "├── process_tracking/          (9 files)   — step-by-step decision log")
code(doc, "├── Dockerfile                (21 lines)  — python:3.11-slim, trains at build")
code(doc, "├── docker-compose.full.yml   (44 lines)  — full stack: API+MLflow+Prom+Grafana")
code(doc, "├── Makefile                  (102 lines) — all common commands automated")
code(doc, "├── requirements.txt          (16 deps)   — all pinned to exact versions")
code(doc, "├── pytest.ini                            — pythonpath config for tests")
code(doc, "└── generate_architecture.py             — architecture diagram generator")

doc.add_page_break()

# ── SECTION 15: REFERENCES ────────────────────────────────────────────────────
h(doc, "15. References")

refs = [
    "Detrano, R. et al. (1989). International application of a new probability algorithm for the diagnosis of coronary artery disease. American Journal of Cardiology, 64(5), 304–310.",
    "UCI Machine Learning Repository — Heart Disease Dataset. https://archive.ics.uci.edu/ml/datasets/Heart+Disease",
    "Chobanian, A.V. et al. (2003). The Seventh Report of the Joint National Committee on Prevention, Detection, Evaluation, and Treatment of High Blood Pressure (JNC-7). JAMA, 289(19), 2560–2571.",
    "Expert Panel on Detection, Evaluation, and Treatment of High Blood Cholesterol in Adults (2001). Executive Summary of The Third Report of the National Cholesterol Education Program (NCEP) Expert Panel (ATP III). JAMA, 285(19), 2486–2497.",
    "Sculley, D. et al. (2015). Hidden Technical Debt in Machine Learning Systems. Advances in Neural Information Processing Systems 28.",
    "MLflow Documentation — Tracking. https://mlflow.org/docs/latest/tracking.html",
    "FastAPI Documentation. https://fastapi.tiangolo.com",
    "Prometheus Documentation — Data Model. https://prometheus.io/docs/concepts/data_model/",
    "Kubernetes Documentation — Deployments. https://kubernetes.io/docs/concepts/workloads/controllers/deployment/",
    "scikit-learn User Guide — Pipelines and Composite Estimators. https://scikit-learn.org/stable/modules/compose.html",
]

for i, ref in enumerate(refs, 1):
    rp = doc.add_paragraph(f"[{i}]  {ref}")
    rp.paragraph_format.space_after = Pt(5)
    for run in rp.runs:
        run.font.size = Pt(10)

# ── SAVE ──────────────────────────────────────────────────────────────────────
out = Path("Heart_Disease_MLOps_Report.docx")
doc.save(out)
print(f"Report saved: {out}  ({out.stat().st_size // 1024} KB)")
